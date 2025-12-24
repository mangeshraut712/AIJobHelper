"""Export service for generating resumes in various formats."""
import logging
import os
from datetime import datetime
from typing import Optional
from docx import Document
from fpdf import FPDF
from schemas import ResumeData
from slugify import slugify

logger = logging.getLogger(__name__)


class ExportService:
    """Service for exporting resumes to various formats."""
    
    @staticmethod
    def get_organized_path(job_title: Optional[str] = None, base_dir: str = "storage/exports") -> str:
        """Creates an organized path: storage/exports/YYYY-MM-DD/job-title/"""
        try:
            date_str = datetime.now().strftime("%Y-%m-%d")
            slug_title = slugify(job_title or "general")
            path = os.path.join(base_dir, date_str, slug_title)
            os.makedirs(path, exist_ok=True)
            return path
        except Exception as e:
            logger.error(f"Failed to create organized path: {str(e)}")
            raise
    
    @staticmethod
    def to_docx(resume: ResumeData, output_path: str) -> str:
        """Export resume to DOCX format."""
        try:
            if not resume.name:
                raise ValueError("Resume name is required")
            
            doc = Document()
            doc.add_heading(resume.name, 0)
            
            # Contact information
            contact_parts = [resume.email]
            if resume.phone:
                contact_parts.append(resume.phone)
            if resume.linkedin:
                contact_parts.append(resume.linkedin)
            
            if contact_parts:
                p = doc.add_paragraph()
                p.add_run(" | ".join(contact_parts)).italic = True
            
            # Summary
            if resume.summary:
                doc.add_heading('Summary', level=1)
                doc.add_paragraph(resume.summary)
            
            # Experience
            if resume.experience:
                doc.add_heading('Experience', level=1)
                for exp in resume.experience:
                    role = exp.get('role', 'N/A')
                    company = exp.get('company', 'N/A')
                    doc.add_heading(f"{role} at {company}", level=2)
                    
                    if exp.get('duration'):
                        doc.add_paragraph(exp['duration']).italic = True
                    
                    if exp.get('description'):
                        doc.add_paragraph(exp['description'])
            
            # Education
            if resume.education:
                doc.add_heading('Education', level=1)
                for edu in resume.education:
                    institution = edu.get('institution', 'N/A')
                    doc.add_heading(institution, level=2)
                    
                    degree = edu.get('degree', '')
                    year = edu.get('graduation_year', '')
                    degree_text = degree
                    if year:
                        degree_text += f" ({year})"
                    doc.add_paragraph(degree_text)
            
            # Skills
            if resume.skills:
                doc.add_heading('Skills', level=1)
                doc.add_paragraph(", ".join(resume.skills))
            
            # Ensure directory exists
            os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else ".", exist_ok=True)
            
            doc.save(output_path)
            logger.info(f"Successfully exported DOCX to: {output_path}")
            return output_path
        except Exception as e:
            logger.error(f"Failed to export DOCX: {str(e)}", exc_info=True)
            raise

    @staticmethod
    def to_pdf(resume: ResumeData, output_path: str) -> str:
        """Export resume to PDF format."""
        try:
            if not resume.name:
                raise ValueError("Resume name is required")
            
            pdf = FPDF()
            pdf.add_page()
            
            # Name
            pdf.set_font("Helvetica", "B", 16)
            pdf.cell(0, 10, resume.name, ln=True, align='C')
            
            # Contact information
            contact_parts = [resume.email]
            if resume.phone:
                contact_parts.append(resume.phone)
            if resume.linkedin:
                contact_parts.append(resume.linkedin)
            
            if contact_parts:
                pdf.set_font("Helvetica", "I", 10)
                pdf.cell(0, 10, " | ".join(contact_parts), ln=True, align='C')
            
            # Summary
            if resume.summary:
                pdf.set_font("Helvetica", "B", 12)
                pdf.cell(0, 10, "Summary", ln=True)
                pdf.set_font("Helvetica", "", 10)
                pdf.multi_cell(0, 5, resume.summary)
            
            # Experience
            if resume.experience:
                pdf.set_font("Helvetica", "B", 12)
                pdf.cell(0, 10, "Experience", ln=True)
                for exp in resume.experience:
                    role = exp.get('role', 'N/A')
                    company = exp.get('company', 'N/A')
                    pdf.set_font("Helvetica", "B", 10)
                    pdf.cell(0, 7, f"{role} at {company}", ln=True)
                    
                    if exp.get('duration'):
                        pdf.set_font("Helvetica", "I", 9)
                        pdf.cell(0, 7, exp['duration'], ln=True)
                    
                    if exp.get('description'):
                        pdf.set_font("Helvetica", "", 10)
                        pdf.multi_cell(0, 5, exp['description'])
            
            # Education
            if resume.education:
                pdf.set_font("Helvetica", "B", 12)
                pdf.cell(0, 10, "Education", ln=True)
                for edu in resume.education:
                    institution = edu.get('institution', 'N/A')
                    degree = edu.get('degree', '')
                    year = edu.get('graduation_year', '')
                    
                    pdf.set_font("Helvetica", "B", 10)
                    pdf.cell(0, 7, institution, ln=True)
                    pdf.set_font("Helvetica", "", 10)
                    degree_text = degree
                    if year:
                        degree_text += f" ({year})"
                    pdf.cell(0, 7, degree_text, ln=True)
            
            # Skills
            if resume.skills:
                pdf.set_font("Helvetica", "B", 12)
                pdf.cell(0, 10, "Skills", ln=True)
                pdf.set_font("Helvetica", "", 10)
                pdf.multi_cell(0, 5, ", ".join(resume.skills))
            
            # Ensure directory exists
            os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else ".", exist_ok=True)
            
            pdf.output(output_path)
            logger.info(f"Successfully exported PDF to: {output_path}")
            return output_path
        except Exception as e:
            logger.error(f"Failed to export PDF: {str(e)}", exc_info=True)
            raise

    @staticmethod
    def to_latex(resume: ResumeData) -> str:
        """Export resume to LaTeX format."""
        try:
            if not resume.name:
                raise ValueError("Resume name is required")
            
            # Escape LaTeX special characters
            def escape_latex(text: str) -> str:
                if not text:
                    return ""
                replacements = {
                    '&': r'\&',
                    '%': r'\%',
                    '$': r'\$',
                    '#': r'\#',
                    '^': r'\textasciicircum{}',
                    '_': r'\_',
                    '{': r'\{',
                    '}': r'\}',
                    '~': r'\textasciitilde{}',
                    '\\': r'\textbackslash{}',
                }
                for char, replacement in replacements.items():
                    text = text.replace(char, replacement)
                return text
            
            name_parts = resume.name.split()
            first_name = escape_latex(name_parts[0]) if name_parts else ""
            last_name = escape_latex(" ".join(name_parts[1:])) if len(name_parts) > 1 else ""
            
            latex = r"""
\documentclass[11pt,a4paper,sans]{moderncv}
\moderncvstyle{casual}
\moderncvcolor{blue}
\name{""" + first_name + "}{" + last_name + r"""}
\address{""" + escape_latex(resume.email) + r"""}
"""
            
            if resume.phone:
                latex += r"\phone[mobile]{" + escape_latex(resume.phone) + r"}\n"
            
            if resume.linkedin:
                latex += r"\social[linkedin]{" + escape_latex(resume.linkedin) + r"}\n"
            
            latex += r"""
\begin{document}
\makecvtitle

"""
            
            # Summary
            if resume.summary:
                latex += r"\section{Summary}" + "\n"
                latex += escape_latex(resume.summary) + "\n\n"
            
            # Experience
            if resume.experience:
                latex += r"\section{Experience}" + "\n"
                for exp in resume.experience:
                    duration = escape_latex(exp.get('duration', ''))
                    role = escape_latex(exp.get('role', ''))
                    company = escape_latex(exp.get('company', ''))
                    description = escape_latex(exp.get('description', ''))
                    latex += f"\\cventry{{{duration}}}{{{role}}}{{{company}}}{{}}{{}}{{{description}}}\n"
                latex += "\n"
            
            # Education
            if resume.education:
                latex += r"\section{Education}" + "\n"
                for edu in resume.education:
                    year = escape_latex(str(edu.get('graduation_year', '')))
                    degree = escape_latex(edu.get('degree', ''))
                    institution = escape_latex(edu.get('institution', ''))
                    latex += f"\\cventry{{{year}}}{{{degree}}}{{{institution}}}{{}}{{}}{{}}\n"
                latex += "\n"
            
            # Skills
            if resume.skills:
                latex += r"\section{Skills}" + "\n"
                skills_text = ", ".join([escape_latex(skill) for skill in resume.skills])
                latex += f"\\cvitem{{Core Skills}}{{{skills_text}}}\n\n"
            
            latex += r"""
\end{document}
"""
            
            logger.info("Successfully generated LaTeX export")
            return latex
        except Exception as e:
            logger.error(f"Failed to export LaTeX: {str(e)}", exc_info=True)
            raise

