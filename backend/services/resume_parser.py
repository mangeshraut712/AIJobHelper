"""Resume parser for extracting text from various file formats."""
import logging
import PyPDF2
import docx
import io
from typing import Optional

logger = logging.getLogger(__name__)


class ResumeParser:
    """Parser for extracting text from resume files."""
    
    @staticmethod
    def extract_text_from_pdf(file_bytes: bytes) -> str:
        """Extract text from PDF file."""
        text = ""
        try:
            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            if len(reader.pages) == 0:
                logger.warning("PDF file has no pages")
                return text
            
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num + 1}: {str(e)}")
                    continue
            
            if not text.strip():
                logger.warning("No text extracted from PDF")
        except PyPDF2.errors.PdfReadError as e:
            logger.error(f"PDF read error: {str(e)}")
            raise ValueError(f"Invalid PDF file: {str(e)}")
        except Exception as e:
            logger.error(f"PDF extraction error: {str(e)}")
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
        
        return text.strip()

    @staticmethod
    def extract_text_from_docx(file_bytes: bytes) -> str:
        """Extract text from DOCX file."""
        text = ""
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        text += row_text + "\n"
            
            if not text.strip():
                logger.warning("No text extracted from DOCX")
        except Exception as e:
            logger.error(f"DOCX extraction error: {str(e)}")
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
        
        return text.strip()

    @classmethod
    def extract_text(cls, file_bytes: bytes, filename: str) -> str:
        """Extract text from file based on extension."""
        if not filename:
            raise ValueError("Filename is required")
        
        filename_lower = filename.lower()
        
        try:
            if filename_lower.endswith(".pdf"):
                return cls.extract_text_from_pdf(file_bytes)
            elif filename_lower.endswith((".docx", ".doc")):
                return cls.extract_text_from_docx(file_bytes)
            elif filename_lower.endswith(".txt"):
                try:
                    return file_bytes.decode("utf-8")
                except UnicodeDecodeError:
                    return file_bytes.decode("latin-1", errors="ignore")
            else:
                return file_bytes.decode("utf-8", errors="ignore")
        except ValueError:
            raise
        except Exception as e:
            logger.error(f"Unexpected error extracting text: {str(e)}")
            raise ValueError(f"Failed to extract text from file: {str(e)}")

