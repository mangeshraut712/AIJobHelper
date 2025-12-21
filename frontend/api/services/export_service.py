from docx import Document
from fpdf import FPDF
from schemas import ResumeData
import os
from datetime import datetime
from slugify import slugify

class ExportService:
    @staticmethod
    def get_organized_path(job_title: str, base_dir: str = "storage/exports") -> str:
        """
        Creates an organized path: storage/exports/YYYY-MM-DD/job-title/
        """
        date_str = datetime.now().strftime("%Y-%m-%d")
        slug_title = slugify(job_title or "general")
        path = os.path.join(base_dir, date_str, slug_title)
        os.makedirs(path, exist_ok=True)
        return path
    @staticmethod
    def to_docx(resume: ResumeData, output_path: str):
        doc = Document()
        doc.add_heading(resume.name, 0)
        
        p = doc.add_paragraph()
        p.add_run(f"{resume.email} | {resume.phone} | {resume.linkedin}").italic = True
        
        doc.add_heading('Summary', level=1)
        doc.add_paragraph(resume.summary)
        
        doc.add_heading('Experience', level=1)
        for exp in resume.experience:
            doc.add_heading(f"{exp['role']} at {exp['company']}", level=2)
            doc.add_paragraph(exp['duration']).italic = True
            doc.add_paragraph(exp['description'])
            
        doc.add_heading('Education', level=1)
        for edu in resume.education:
            doc.add_heading(edu['institution'], level=2)
            doc.add_paragraph(f"{edu['degree']} ({edu['graduation_year']})")
            
        doc.add_heading('Skills', level=1)
        doc.add_paragraph(", ".join(resume.skills))
        
        doc.save(output_path)
        return output_path

    @staticmethod
    def to_pdf(resume: ResumeData, output_path: str):
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 10, resume.name, ln=True, align='C')
        
        pdf.set_font("Helvetica", "I", 10)
        pdf.cell(0, 10, f"{resume.email} | {resume.phone} | {resume.linkedin}", ln=True, align='C')
        
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 10, "Summary", ln=True)
        pdf.set_font("Helvetica", "", 10)
        pdf.multi_cell(0, 5, resume.summary)
        
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 10, "Experience", ln=True)
        for exp in resume.experience:
            pdf.set_font("Helvetica", "B", 10)
            pdf.cell(0, 7, f"{exp['role']} at {exp['company']}", ln=True)
            pdf.set_font("Helvetica", "I", 9)
            pdf.cell(0, 7, exp['duration'], ln=True)
            pdf.set_font("Helvetica", "", 10)
            pdf.multi_cell(0, 5, exp['description'])
            
        pdf.output(output_path)
        return output_path

    @staticmethod
    def to_latex(resume: ResumeData) -> str:
        # Simple LaTeX template
        latex = r"""
\documentclass[11pt,a4paper,sans]{moderncv}
\moderncvstyle{casual}
\moderncvcolor{blue}
\name{""" + resume.name.split()[0] + "}{" + " ".join(resume.name.split()[1:]) + r"""}
\address{""" + resume.email + r"""}
\phone[mobile]{""" + (resume.phone or "") + r"""}
\social[linkedin]{""" + (resume.linkedin or "") + r"""}

\begin{document}
\makecvtitle

\section{Summary}
""" + resume.summary + r"""

\section{Experience}
"""
        for exp in resume.experience:
            latex += r"\cventry{" + exp['duration'] + "}{" + exp['role'] + "}{" + exp['company'] + "}{}{}{" + exp['description'] + "}\n"
            
        latex += r"""
\section{Education}
"""
        for edu in resume.education:
            latex += r"\cventry{" + str(edu['graduation_year']) + "}{" + edu['degree'] + "}{" + edu['institution'] + "}{}{}{}\n"
            
        latex += r"""
\section{Skills}
\cvitem{Core Skills}{""" + ", ".join(resume.skills) + r"""}

\end{document}
"""
        return latex
